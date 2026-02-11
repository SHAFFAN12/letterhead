from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Constants
NAVY_BLUE = RGBColor(0x05, 0x11, 0x3b) # #05113b
GOLD = RGBColor(0xc5, 0xa0, 0x59)      # #c5a059
LOGO_PATH = 'logo.png'
QR_PATH = 'qr_code.png'

def set_margins(doc, top, bottom, left, right):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def create_letterhead():
    doc = Document()
    set_margins(doc, 2.5, 2.5, 2.5, 2.5)

    # Header
    header = doc.sections[0].header
    htt = header.add_table(rows=1, cols=2, width=Inches(6))
    htt.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Logo Cell
    cell_logo = htt.cell(0, 0)
    if os.path.exists(LOGO_PATH):
        try:
            p = cell_logo.paragraphs[0]
            run = p.add_run()
            run.add_picture(LOGO_PATH, height=Cm(2.5))
        except Exception as e:
            cell_logo.text = "HORIZON OOH"

    # Brand Text Cell
    cell_text = htt.cell(0, 1)
    p = cell_text.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("HORIZON OOH\n")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY_BLUE
    run = p.add_run("Your Vision Zone")
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD

    # Contact Info in Header
    p = cell_text.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("+1 234 567 8900 | info@horizonooh.com | www.horizon-ooh.com")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Gold Bar under header
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(20)
    border_bottom = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24') # 3pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C5A059')
    border_bottom.append(bottom)
    p._p.get_or_add_pPr().append(border_bottom)

    # Body Content Placeholder
    p = doc.add_paragraph("Date: May 15, 2024")
    doc.add_paragraph("To,\nRecipient Name\nDesignation\nCompany Name\nAddress Line 1")
    doc.add_paragraph("Subject: Your Vision Zone Proposal")
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.")
    
    # Footer
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HORIZON OOH - Your Vision Zone")
    run.bold = True
    run.font.color.rgb = NAVY_BLUE
    p = footer.add_paragraph("123 Any Street, Suite 100, Your City, ST 2345 | www.horizon-ooh.com")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)

    doc.save('Letterhead.docx')
    print("Letterhead.docx created")

def create_invoice():
    doc = Document()
    set_margins(doc, 2.0, 2.0, 2.0, 2.0)

    # Header (Simplified)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("INVOICE\n")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY_BLUE
    run = p.add_run("Invoice #: INV-2024-001\nDate: May 15, 2024")

    # Client Info
    doc.add_paragraph("Bill To:").runs[0].bold = True
    doc.add_paragraph("Client Name\nClient Company\nAddress")

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Description', 'Qty', 'Unit Price', 'Total']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Items
    items = [
        ("Billboard Advertising", "1", "$5,000.00", "$5,000.00"),
        ("Production & Installation", "1", "$1,200.00", "$1,200.00"),
        ("Design Consultation", "5", "$100.00", "$500.00")
    ]
    for item in items:
        row_cells = table.add_row().cells
        for i, val in enumerate(item):
            row_cells[i].text = val

    # Calculation
    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Subtotal: $6,700.00\n").bold = True
    p.add_run("Tax (10%): $670.00\n")
    run = p.add_run("Total: $7,370.00")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY_BLUE

    doc.save('Invoice.docx')
    print("Invoice.docx created")

def create_visiting_card():
    # Attempting to create a small doc size
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(3.5)
    section.page_height = Inches(2.0)
    set_margins(doc, 0.2, 0.2, 0.2, 0.2)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Left: Logo
    cell_left = table.cell(0, 0)
    cell_left.width = Inches(1.5)
    if os.path.exists(LOGO_PATH):
        try:
            p = cell_left.paragraphs[0]
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(1.0))
        except: pass
    
    # Right: Details
    cell_right = table.cell(0, 1)
    cell_right.width = Inches(1.8)
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Ghufran Muhammed\n")
    run.bold = True
    run.font.size = Pt(12)
    
    run = p.add_run("CEO\n\n")
    run.italic = True
    run.font.size = Pt(10)

    run = p.add_run("03091329000\ninfo@horizonooh.com")
    run.font.size = Pt(8)

    doc.save('Visiting_Card.docx')
    print("Visiting_Card.docx created")

def create_envelope():
    # DL Size: 220mm x 110mm = 8.66in x 4.33in
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(22)
    section.page_height = Cm(11)
    set_margins(doc, 1.0, 1.0, 1.0, 1.0)

    # Logo
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, height=Cm(1.5))
    
    doc.add_paragraph("HORIZON OOH - Your Vision Zone\n123 Any Street, Suite 100, Your City, ST 2345")

    # To Address
    p = doc.add_paragraph("\n\n\nTo:\n__________________________\n__________________________")
    p.paragraph_format.left_indent = Inches(3.0)

    doc.save('Envelope.docx')
    print("Envelope.docx created")

if __name__ == "__main__":
    create_letterhead()
    create_invoice()
    create_visiting_card()
    create_envelope()
